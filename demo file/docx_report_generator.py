"""
DOCX 형식 보고서 생성기
python-docx를 사용하여 전문적인 Word 문서 생성
"""

from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx --break-system-packages")


class DocxReportGenerator:
    """DOCX 형식 보고서 생성기"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """문서 스타일 설정"""
        # 기본 폰트 설정
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Noto Sans KR'
        font.size = Pt(10)
        
        # 제목 스타일
        heading1 = self.doc.styles['Heading 1']
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(102, 126, 234)  # #667eea
    
    def add_cover_page(self, company: str, country: str):
        """표지 페이지 추가"""
        # 제목
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{company} × {country}\n")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 126, 234)
        
        # 부제
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("다국적 시장 진출 전략 보고서")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(102, 102, 102)
        
        self.doc.add_paragraph()  # 공백
        
        # 날짜
        date = self.doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date.add_run(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        self.doc.add_page_break()
    
    def add_executive_summary(self, decision_data: Dict, regulatory_data: Dict, partner_count: int):
        """Executive Summary 섹션"""
        self.doc.add_heading('📋 Executive Summary', 1)
        
        # 최종 의사결정
        self.doc.add_heading('최종 의사결정', 2)
        decision = self.doc.add_paragraph()
        run = decision.add_run(f"결정: {decision_data.get('decision', '보류')}")
        run.bold = True
        run.font.size = Pt(14)
        
        score = self.doc.add_paragraph()
        run = score.add_run(f"최종 점수: {decision_data.get('final_score', 0)}점")
        run.bold = True
        run.font.size = Pt(14)
        
        rationale = self.doc.add_paragraph(decision_data.get('rationale', '분석 중...'))
        rationale.style = 'Intense Quote'
        
        # 주요 지표
        self.doc.add_heading('주요 지표', 2)
        
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = [
            ('규제 커버리지', f"{regulatory_data.get('coverage_score', 0):.0%}"),
            ('규제 리스크', regulatory_data.get('risk_badge', '보통')),
            ('파트너 후보', f"{partner_count}개사")
        ]
        
        for i, (label, value) in enumerate(cells):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    def add_market_research(self, market_summary: Dict, image_path: Optional[str] = None):
        """시장 조사 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('🌍 시장 조사', 1)
        
        # 핵심 지표
        self.doc.add_heading('핵심 지표', 2)
        
        metrics = market_summary.get('metrics', {})
        if metrics:
            table = self.doc.add_table(rows=len(metrics), cols=2)
            table.style = 'Light List Accent 1'
            
            for i, (key, value) in enumerate(metrics.items()):
                row = table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = str(value)
                row.cells[1].paragraphs[0].runs[0].bold = True
        
        # Why Now
        self.doc.add_heading('Why Now?', 2)
        why_now = self.doc.add_paragraph(market_summary.get('why_now', '시장 분석 중...'))
        why_now.style = 'Intense Quote'
        
        # 이미지 추가
        if image_path and Path(image_path).exists():
            self.doc.add_heading('수요 히트맵', 2)
            self.doc.add_picture(image_path, width=Inches(6))
    
    def add_regulatory_review(self, regulatory_data: Dict, image_path: Optional[str] = None):
        """규제 검토 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('✅ 규제 검토', 1)
        
        # 요약
        summary = self.doc.add_paragraph()
        summary.add_run(f"커버리지: {regulatory_data.get('coverage_score', 0):.0%} | ")
        run = summary.add_run(f"리스크: {regulatory_data.get('risk_badge', '보통')}")
        run.bold = True
        
        # 체크리스트 테이블
        self.doc.add_heading('규제 체크리스트', 2)
        
        checklist = regulatory_data.get('checklist', [])
        if checklist:
            table = self.doc.add_table(rows=len(checklist) + 1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # 헤더
            headers = ['항목', '등급', '상태', '근거/출처']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # 데이터
            for i, item in enumerate(checklist, 1):
                row = table.rows[i]
                row.cells[0].text = item.get('item', '-')
                row.cells[1].text = item.get('grade', '-')
                row.cells[2].text = item.get('state', '-')
                row.cells[3].text = item.get('evidence', '-')
        
        # 통관 플로우 이미지
        if image_path and Path(image_path).exists():
            self.doc.add_heading('통관·관세 플로우', 2)
            self.doc.add_picture(image_path, width=Inches(6))
    
    def add_competitor_analysis(
        self,
        competitor_data: Dict,
        density_map_path: Optional[str] = None,
        positioning_path: Optional[str] = None
    ):
        """경쟁사 분석 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('🗺️ 경쟁사 분석', 1)
        
        # 화이트스페이스
        self.doc.add_heading('화이트스페이스 (기회 영역)', 2)
        
        gaps = competitor_data.get('whitespace_gaps', [])
        for i, gap in enumerate(gaps, 1):
            p = self.doc.add_paragraph(gap, style='List Number')
        
        # 경쟁 밀집 지도
        if density_map_path and Path(density_map_path).exists():
            self.doc.add_heading('경쟁 밀집 지도', 2)
            self.doc.add_picture(density_map_path, width=Inches(6))
        
        # 포지셔닝 매트릭스
        if positioning_path and Path(positioning_path).exists():
            self.doc.add_heading('포지셔닝 매트릭스', 2)
            self.doc.add_picture(positioning_path, width=Inches(6))
    
    def add_partner_discovery(self, partner_data: Dict, map_path: Optional[str] = None):
        """파트너 발굴 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('🤝 파트너 발굴', 1)
        
        # 파트너 후보 테이블
        candidates = partner_data.get('candidates', [])
        if candidates:
            self.doc.add_heading(f'파트너 후보 ({len(candidates)}개사)', 2)
            
            table = self.doc.add_table(rows=len(candidates) + 1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # 헤더
            headers = ['업체명', '역할', '선정 이유', '대안']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # 데이터
            for i, partner in enumerate(candidates, 1):
                row = table.rows[i]
                row.cells[0].text = partner.get('name', '-')
                row.cells[1].text = partner.get('role', '-')
                row.cells[2].text = partner.get('rationale', '-')
                row.cells[3].text = partner.get('alternative', '-')
        
        # 파트너 맵
        if map_path and Path(map_path).exists():
            self.doc.add_heading('파트너 연결 맵', 2)
            self.doc.add_picture(map_path, width=Inches(6))
        
        # PoC 제안서
        poc = partner_data.get('poc_proposal', '')
        if poc:
            self.doc.add_heading('PoC 제안서 초안', 2)
            self.doc.add_paragraph(poc, style='Intense Quote')
    
    def add_risk_scenarios(self, risk_data: Dict):
        """리스크 시나리오 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('⚠️ 리스크 시나리오', 1)
        
        # 리스크 레지스터
        self.doc.add_heading('리스크 레지스터', 2)
        
        risks = risk_data.get('register', [])
        if risks:
            table = self.doc.add_table(rows=len(risks) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # 헤더
            headers = ['리스크', '가능성', '영향', '초기 징후', '완화 전략']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # 데이터
            for i, risk in enumerate(risks, 1):
                row = table.rows[i]
                row.cells[0].text = risk.get('risk', '-')
                row.cells[1].text = risk.get('likelihood', '-')
                row.cells[2].text = risk.get('impact', '-')
                row.cells[3].text = risk.get('early_sign', '-')
                row.cells[4].text = risk.get('mitigation', '-')
        
        # 시나리오
        self.doc.add_heading('시나리오별 KPI 범위', 2)
        
        scenarios = risk_data.get('scenarios', {})
        for scenario_name, kpis in scenarios.items():
            scenario_label = {
                'pessimistic': '😰 비관 시나리오',
                'base': '😐 기준 시나리오',
                'optimistic': '😄 낙관 시나리오'
            }.get(scenario_name, scenario_name)
            
            p = self.doc.add_paragraph()
            run = p.add_run(scenario_label)
            run.bold = True
            run.font.size = Pt(12)
            
            for key, value in kpis.items():
                self.doc.add_paragraph(f"{key}: {value}", style='List Bullet')
    
    def add_decision_scorecard(self, decision_data: Dict):
        """의사결정 스코어카드 섹션"""
        self.doc.add_page_break()
        self.doc.add_heading('🎯 의사결정 스코어카드', 1)
        
        # 스코어카드 테이블
        scorecard = decision_data.get('scorecard', {})
        if scorecard:
            table = self.doc.add_table(rows=len(scorecard), cols=2)
            table.style = 'Medium Shading 1 Accent 1'
            
            for i, (category, score) in enumerate(scorecard.items()):
                row = table.rows[i]
                row.cells[0].text = category
                row.cells[1].text = str(score)
                row.cells[1].paragraphs[0].runs[0].bold = True
        
        self.doc.add_paragraph()  # 공백
        
        # 최종 결론
        self.doc.add_heading('최종 결론', 2)
        
        conclusion = self.doc.add_paragraph()
        run = conclusion.add_run(f"{decision_data.get('decision', '보류')}\n")
        run.bold = True
        run.font.size = Pt(16)
        
        self.doc.add_paragraph(decision_data.get('rationale', '분석 중...'))
        
        # Next Steps
        next_steps = decision_data.get('next_steps', [])
        if next_steps:
            self.doc.add_heading('📌 Next Steps', 2)
            for step in next_steps:
                self.doc.add_paragraph(step, style='List Number')
    
    def save(self, filepath: str):
        """문서 저장"""
        self.doc.save(filepath)
        print(f"DOCX 보고서 저장 완료: {filepath}")
    
    @classmethod
    def generate_full_report(
        cls,
        company: str,
        country: str,
        market_summary: Dict,
        regulatory_data: Dict,
        competitor_data: Dict,
        partner_data: Dict,
        risk_data: Dict,
        decision_data: Dict,
        visualizations: Dict[str, str],
        output_path: str
    ):
        """전체 보고서 생성 (원스텝)"""
        generator = cls()
        
        # 표지
        generator.add_cover_page(company, country)
        
        # Executive Summary
        generator.add_executive_summary(
            decision_data,
            regulatory_data,
            len(partner_data.get('candidates', []))
        )
        
        # 시장 조사
        generator.add_market_research(
            market_summary,
            visualizations.get('demand_heatmap')
        )
        
        # 규제 검토
        generator.add_regulatory_review(
            regulatory_data,
            visualizations.get('customs_flow')
        )
        
        # 경쟁사 분석
        generator.add_competitor_analysis(
            competitor_data,
            visualizations.get('competition_density'),
            visualizations.get('positioning')
        )
        
        # 파트너 발굴
        generator.add_partner_discovery(
            partner_data,
            visualizations.get('partner_map')
        )
        
        # 리스크 시나리오
        generator.add_risk_scenarios(risk_data)
        
        # 의사결정 스코어카드
        generator.add_decision_scorecard(decision_data)
        
        # 저장
        generator.save(output_path)
        
        return output_path


# 사용 예시
if __name__ == "__main__":
    sample_data = {
        "company": "스타트업A",
        "country": "베트남",
        "market_summary": {
            "metrics": {
                "TAM": "$2.5B",
                "성장률": "15% YoY",
                "택배량": "500M건/년"
            },
            "why_now": "이커머스 급성장과 중산층 확대"
        },
        "regulatory_data": {
            "coverage_score": 0.76,
            "risk_badge": "보통",
            "checklist": [
                {"item": "택배업 등록", "grade": "MUST", "state": "PASS", "evidence": "국토부 고시"}
            ]
        },
        "competitor_data": {
            "whitespace_gaps": ["농촌 지역 당일배송", "콜드체인 물류"]
        },
        "partner_data": {
            "candidates": [
                {"name": "VietPost", "role": "라스트마일", "rationale": "전국 네트워크", "alternative": "GHN"}
            ],
            "poc_proposal": "3개월 시범 운영"
        },
        "risk_data": {
            "register": [
                {"risk": "규제 변경", "likelihood": "보통", "impact": "높음", "early_sign": "정부 발표", "mitigation": "로비"}
            ],
            "scenarios": {
                "base": {"월 거래량": "30K건"}
            }
        },
        "decision_data": {
            "decision": "진행(추천)",
            "rationale": "시장 성장성 높음",
            "final_score": 72,
            "scorecard": {"시장 매력도": 85},
            "next_steps": ["DPA 협약"]
        },
        "visualizations": {}
    }
    
    output = "/tmp/test_report.docx"
    DocxReportGenerator.generate_full_report(**sample_data, output_path=output)
    print(f"테스트 완료: {output}")

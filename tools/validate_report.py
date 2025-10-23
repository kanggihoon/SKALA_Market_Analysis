import sys
import os
import re
from docx import Document


def load_doc(path):
    try:
        return Document(path)
    except Exception as e:
        print(f"FAIL: cannot open report: {e}")
        sys.exit(1)


def doc_text(doc):
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                texts.append(c.text)
    return "\n".join(texts)


def assert_contains(txt, needle, msg):
    if needle not in txt:
        print(f"FAIL: missing '{needle}' ({msg})")
        sys.exit(1)


def main():
    report = os.path.join('outputs', 'Final_Report.docx')
    if not os.path.exists(report):
        # try timestamped fallback
        candidates = [p for p in os.listdir('outputs') if p.startswith('Final_Report_') and p.endswith('.docx')]
        if candidates:
            candidates.sort()
            report = os.path.join('outputs', candidates[-1])
        else:
            print('FAIL: report not found')
            sys.exit(1)

    doc = load_doc(report)
    txt = doc_text(doc)

    # 1) headers
    for case in [
        'ShipBob × KR',
        'ShipBob × JP',
        'Locus.sh × US',
        'Ninja Van × JP',
    ]:
        assert_contains(txt, case, 'case header')

    # section headings
    for sec in ['## Market','## Regulation','## Competition','## GTM','## Partners','## Risks','## Decision Scorecard']:
        assert_contains(txt, sec, 'section headings')

    # 2) Coverage pattern at least 4 and not repeated 3+ times same value
    covs = re.findall(r"Coverage: (\d+)%", txt)
    if len(covs) < 4:
        print('FAIL: coverage occurrences < 4')
        sys.exit(1)
    # repetition check
    from collections import Counter
    counts = Counter(covs)
    if any(v >= 3 for v in counts.values()):
        print('FAIL: same Coverage value repeated >=3 times')
        sys.exit(1)

    # 3) image paths existence per case
    def _check_img(case_name, iso):
        base = os.path.join('outputs', f"{case_name}_{iso}")
        ok = True
        ms = [p for p in os.listdir(base) if p.startswith('01_market_summary_')]
        cf = [p for p in os.listdir(base) if p.startswith('02_customs_flow_')]
        mp = [p for p in os.listdir(base) if p.startswith('map_')]
        if not ms or not cf or not mp:
            return False
        return True

    if not (_check_img('ShipBob','KR') and _check_img('ShipBob','JP') and _check_img('Locus.sh','US') and _check_img('Ninja Van','JP')):
        print('FAIL: required image paths missing in one or more cases')
        sys.exit(1)

    # 4) Segment score equality
    seg_scores = re.findall(r"\|\s*(high|mid|low)\s*\|\s*([0-9]+\.[0-9])\s*\|", txt, re.I)
    if seg_scores:
        scores_only = [s for _, s in seg_scores]
        if len(set(scores_only)) == 1:
            print('FAIL: segment scores are identical')
            sys.exit(1)

    # 5) MUST-FAIL implies HOLD
    if 'MUST-FAIL' in txt and 'Decision: RECOMMEND' in txt:
        print('FAIL: MUST-FAIL present but decision is RECOMMEND')
        sys.exit(1)

    print('PASS: report ok')


if __name__ == '__main__':
    main()


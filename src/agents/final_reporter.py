import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from loguru import logger
from .narrative import generate_texts


def run(state, meta, out_dir: str):
    """Aggregate all company x country outputs into a single Word report (.docx)."""
    os.makedirs(out_dir, exist_ok=True)
    keep_all = str(os.getenv('KEEP_ALL_FINALS','0')).lower() in ('1','true','yes')
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    path = os.path.join(out_dir, f"Final_Report_{ts}.docx") if keep_all else os.path.join(out_dir, "Final_Report.docx")
    doc = Document()
    # compact spacing to reduce blank areas
    try:
        base = doc.styles['Normal']
        base.font.size = Pt(10)
        base.paragraph_format.space_after = Pt(2)
        base.paragraph_format.space_before = Pt(0)
    except Exception:
        pass
    doc.add_heading('Market Entry Strategy Report', level=1)

    for company in meta.get("companies", []):
        name = company.get("name")
        for country in company.get("target_countries", []):
            section = os.path.join(out_dir, f"{name}_{country}")
            doc.add_heading(f"{name} x {country}", level=2)
            doc.add_paragraph("1) Executive / Decision")

            # Load case state
            case = None
            case_json = os.path.join(section, "case_state.json")
            if os.path.exists(case_json):
                try:
                    import json
                    with open(case_json, "r", encoding="utf-8") as f:
                        case = json.load(f)
                except Exception:
                    case = None

            cov = (case or {}).get('coverage') if case else None
            tbd = (case or {}).get('tbd_ratio') if case else None
            badge = (case or {}).get('risk_badge') if case else ''
            dec = (case or {}).get('decision', {}) if case else {}
            sc0 = (dec.get('scorecard', {}) or {})
            final0 = sc0.get('final')

            # KPI badges table
            kpi = doc.add_table(rows=1, cols=4)
            kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = kpi.rows[0].cells
            hdr[0].text = f"Decision: {dec.get('status','')}"
            hdr[1].text = f"Coverage: {round((cov or 0)*100)}%"
            hdr[2].text = f"TBD: {round((tbd or 0)*100)}%"
            hdr[3].text = f"Risk: {badge or ''}"

            for i, cell in enumerate(hdr):
                if i == 0:
                    col_hex = 'E2E8F0'
                elif i == 3 and badge == 'High':
                    col_hex = 'FEE2E2'
                elif i == 3 and badge == 'Medium':
                    col_hex = 'FEF3C7'
                elif i == 3 and badge == 'Low':
                    col_hex = 'DCFCE7'
                else:
                    col_hex = 'F1F5F9'
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), col_hex)
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

            # LLM executive narrative (optional)
            use_llm = str(os.getenv('USE_LLM_NARRATIVE','0')).lower() in ('1','true','yes')
            llm_texts = generate_texts(case) if (use_llm and case) else {}
            if llm_texts.get('exec'):
                doc.add_paragraph(llm_texts['exec'])

            # Helpers for images (standardized filenames)
            def _img(path: str, width_in: float = 6.0):
                try:
                    if os.path.exists(path):
                        doc.add_picture(path, width=Inches(width_in))
                        return True
                except Exception:
                    pass
                return False

            std_market = os.path.join(section, f"01_market_summary_{name}_{country}.png")
            std_customs = os.path.join(section, f"02_customs_flow_{name}_{country}.png")
            std_heatmap = os.path.join(section, f"03_competition_heatmap_{name}_{country}.png")
            std_partner = os.path.join(section, f"04_partner_map_{name}_{country}.png")
            std_map = os.path.join(section, f"map_{name}_{country}.png")

            # 2) Market
            doc.add_paragraph("2) Market")
            if llm_texts.get('market'):
                doc.add_paragraph(llm_texts['market'])
            if case:
                m = case.get("market", {})
                why = m.get('why_now','')
                metrics = m.get('metrics', {}) or {}
                tam = metrics.get('TAM'); cagr = metrics.get('CAGR'); pen = metrics.get('Ecom Penetration')
                infra = metrics.get('Infra Score'); ship = metrics.get('Avg Ship Cost')
                market_txt = (
                    f"TAM {tam}, CAGR {cagr}, penetration {pen}, infra {infra}, avg ship cost {ship}. "
                    f"Why Now: {why}"
                )
                doc.add_paragraph(market_txt)
            # Market image directly under narrative
            _img(std_market, width_in=6.0)

            # 3) Regulation
            doc.add_paragraph("3) Regulation")
            if llm_texts.get('regulation'):
                doc.add_paragraph(llm_texts['regulation'])
            if case:
                cov_pct = round((cov or 0)*100)
                tbd_pct = round((tbd or 0)*100)
                blocker_txt = 'Yes' if (sc0.get('blocker') or False) else 'No'
                reg_line = f"Coverage {cov_pct}%, TBD {tbd_pct}%, MUST violation {blocker_txt}."
                doc.add_paragraph(reg_line)
            # Customs flow image
            _img(std_customs, width_in=6.0)

            # 4) Competition
            comp = case.get("competition", {}) if case else {}
            ws = comp.get("whitespaces", []) if comp else []
            if ws:
                doc.add_paragraph("4) Competition")
                if llm_texts.get('competition'):
                    doc.add_paragraph(llm_texts['competition'])
                doc.add_paragraph(f"Whitespaces {len(ws)}:")
                for w in ws:
                    doc.add_paragraph(str(w), style="List Bullet")
            ents = comp.get('entities', []) if comp else []
            if ents:
                t = doc.add_table(rows=1, cols=3)
                t.rows[0].cells[0].text = 'Competitor'
                t.rows[0].cells[1].text = 'Category'
                t.rows[0].cells[2].text = 'Homepage'
                for e in ents[:12]:
                    row = t.add_row().cells
                    row[0].text = e.get('name','')
                    row[1].text = e.get('category','')
                    row[2].text = e.get('homepage','')
            # Competition images (heatmap + base map) side-by-side to reduce vertical whitespace
            if os.path.exists(std_heatmap) or os.path.exists(std_map):
                tbl = doc.add_table(rows=1, cols=2)
                cells = tbl.rows[0].cells
                try:
                    if os.path.exists(std_heatmap):
                        run = cells[0].paragraphs[0].add_run()
                        run.add_picture(std_heatmap, width=Inches(3.15))
                    if os.path.exists(std_map):
                        run = cells[1].paragraphs[0].add_run()
                        run.add_picture(std_map, width=Inches(3.15))
                except Exception:
                    # Fallback to stacked if table approach fails
                    _img(std_heatmap, width_in=6.0)
                    _img(std_map, width_in=6.0)

            # 5) GTM
            gtm = case.get("gtm", {}) if case else {}
            table_rows = gtm.get("table", []) if gtm else []
            if table_rows:
                doc.add_paragraph("5) GTM (High/Mid/Low → selected)")
                if llm_texts.get('gtm'):
                    doc.add_paragraph(llm_texts['gtm'])
                t = doc.add_table(rows=1, cols=4)
                hdr = t.rows[0].cells
                hdr[0].text = "Segment"; hdr[1].text = "Score"; hdr[2].text = "ICP"; hdr[3].text = "Offer"
                for r in table_rows:
                    row = t.add_row().cells
                    row[0].text = str(r.get("segment",""))
                    row[1].text = str(r.get("score",""))
                    row[2].text = str(r.get("icp",""))
                    row[3].text = str(r.get("offer",""))

            # 6) Partners
            partners = case.get("partners", []) if case else []
            if partners:
                doc.add_paragraph("6) Partners")
                if llm_texts.get('partners'):
                    doc.add_paragraph(llm_texts['partners'])
                doc.add_paragraph(f"{len(partners)} partner candidates.")
                for p in partners:
                    doc.add_paragraph(f"{p.get('name','')} ({p.get('role','')}) · priority={p.get('priority','')}", style="List Bullet")
            # Partner map image
            _img(std_partner, width_in=6.0)

            # 7) Risks
            risks = case.get("risks", []) if case else []
            if risks:
                doc.add_paragraph("7) Risks")
                if llm_texts.get('risks'):
                    doc.add_paragraph(llm_texts['risks'])
                for r in risks:
                    doc.add_paragraph(f"{r.get('risk','')} · prob={r.get('prob','')} · impact={r.get('impact','')} · mitigation={r.get('mitigation','')} (trigger: {r.get('trigger','')})", style="List Bullet")

            # 8) Decision Scorecard
            if sc0:
                doc.add_paragraph("8) Decision Scorecard")
                t2 = doc.add_table(rows=0, cols=2)
                for k in ["base","cov","tbd_ratio","competition_high","partners","final"]:
                    row = t2.add_row().cells
                    row[0].text = k
                    row[1].text = str(sc0.get(k, ""))

            # 9) Overall
            status = dec.get('status','')
            overall = f"9) Overall: recommend='{status}'."
            if llm_texts.get('overall'):
                doc.add_paragraph(llm_texts['overall'])
            else:
                doc.add_paragraph(overall)

            # 10) 30/60/90 + Evidence (keep concise; images already placed near narratives)
            if any(llm_texts.get(k) for k in ('plan_30','plan_60','plan_90')):
                doc.add_paragraph("10) 30/60/90 plan")
                if llm_texts.get('plan_30'):
                    doc.add_paragraph(llm_texts['plan_30'])
                if llm_texts.get('plan_60'):
                    doc.add_paragraph(llm_texts['plan_60'])
                if llm_texts.get('plan_90'):
                    doc.add_paragraph(llm_texts['plan_90'])

            # Keep pagination minimal: optional page break via env
            if str(os.getenv('PAGE_BREAK_BETWEEN_CASES','0')).lower() in ('1','true','yes'):
                doc.add_page_break()

    # Save file
    try:
        doc.save(path)
    except PermissionError:
        ts_fallback = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt = os.path.join(out_dir, f"Final_Report_{ts_fallback}.docx")
        logger.warning("Final_Report locked ({}). Saving as {}", path, alt)
        doc.save(alt)
        path = alt

    if not keep_all:
        try:
            for f in os.listdir(out_dir):
                if f.startswith("Final_Report_") and f.endswith('.docx') and f != os.path.basename(path):
                    try:
                        os.remove(os.path.join(out_dir, f))
                    except Exception:
                        pass
        except Exception:
            pass
    logger.info("Final report written: {}", path)
    return path

